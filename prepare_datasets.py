import copy
import os
from abc import ABCMeta, abstractmethod
from pathlib import Path
from typing import Tuple
import pandas as pd

from scipy.io import loadmat
from File_Management.load_save_Func import *
from File_Management.path_and_file_management_Func import *
from Ploting.fast_plot_Func import *
from TimeSeries_Class import TimeSeries, UnivariateTimeSeries, merge_two_time_series_df
from Time_Processing.datetime_utils import DatetimeOnehotORCircularEncoder
from project_utils import project_path_
import copy
import getpass
from dateutil import tz
from pandas import DataFrame
from Time_Processing.datetime_utils import DatetimeOnehotORCircularEncoder

try:
    from nilmtk import DataSet, MeterGroup
    from nilmtk.dataset_converters.refit.convert_refit import convert_refit
    from pyorbital.moon_phase import moon_phase
    from workalendar.america import Canada

except ModuleNotFoundError:
    pass

try:
    import torch
    from Writting import docx_document_template_to_collect_figures
    from torch.utils.data import Dataset as TorchDataSet
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_BREAK
    from workalendar.america import Canada
    from pyorbital.moon_phase import moon_phase


    class NILMTorchDataset(TorchDataSet):
        """
        专门针对PyTorch的模型的Dataset
        """
        __slots__ = ('data', 'transformed_data', 'sequence_length', 'over_lapping', 'country')

        def __init__(self, data: pd.DataFrame,
                     *, country,
                     sequence_length: int,
                     over_lapping: Union[bool, int] = False,
                     transform_args_file_path: Path):
            """
            :param data 所有数据，包括x和y，形如下面的一个pd.DataFrame
                               temperature  solar irradiation   precipitation   air density   mains_var   appliance_var
            time_stamp_index
            .
            .
            .
            TODO: solar
            transform_args: 形如下面的一个pd.DataFrame
                      temperature  solar irradiation   precipitation   air density   mains_var   appliance_var

            minimum
            maximum
            :param transform_args_file_path
            """
            self.data = data  # type: pd.DataFrame
            self.sequence_length = sequence_length  # type: int
            self.over_lapping = over_lapping
            self.country = country
            # 最后进行transform
            self.transformed_data = self._transform(
                transform_args_file_path
            )  # type: Tuple[torch.tensor, torch.tensor, torch.tensor]

        def __len__(self):
            if self.over_lapping:
                return self.data.__len__() - self.sequence_length + 1
            else:
                return int(self.data.__len__() / self.sequence_length)

        def __getitem__(self, index: int):
            # 决定索引的位置
            if self.over_lapping:
                index_slice = slice(index, index + self.sequence_length)
            else:
                index_slice = slice(index * self.sequence_length, (index + 1) * self.sequence_length)
            data_x = self.transformed_data[0][index_slice]  # type: torch.tensor
            data_y = self.transformed_data[1][index_slice]  # type: torch.tensor
            return (data_x,
                    data_y,
                    self.transformed_data[2][index_slice],
                    self.transformed_data[2][index_slice])

        def get_raw_data(self, index):
            if self.over_lapping:
                index_slice = slice(index, index + self.sequence_length)
            else:
                index_slice = slice(index * self.sequence_length, (index + 1) * self.sequence_length)
            return self.transformed_data[2][index_slice], self.transformed_data[2][index_slice]

        def _transform(self, transform_args_file_path, mode='NILM') -> Tuple[torch.tensor, torch.tensor, pd.DataFrame]:
            # WARNING: full_data_df.index is not used! because stupid Pandas force to have DST transformation,
            # causing two 1.00 am problem, despite that the original datetime has already been taken care of!
            # full_data_df['index_index'] = range(full_data_df.shape[0]) # This line is for debug
            # index_source = full_data_df.index # This line is to use full_data_df.index
            index_source = pd.DatetimeIndex(pd.date_range(
                start=datetime.datetime(year=self.data.index[0].year,
                                        month=self.data.index[0].month,
                                        day=self.data.index[0].day,
                                        hour=self.data.index[0].hour,
                                        minute=self.data.index[0].minute),
                end=datetime.datetime(year=self.data.index[-1].year,
                                      month=self.data.index[-1].month,
                                      day=self.data.index[-1].day,
                                      hour=self.data.index[-1].hour,
                                      minute=self.data.index[-1].minute),
                freq=self.data.index.freq)
            )
            # 时间变量
            datetime_onehot_or_circular_encoder_1 = DatetimeOnehotORCircularEncoder(to_encoding_args=('holiday',
                                                                                                      'summer_time'),
                                                                                    mode='circular')
            time_var_transformed_1 = datetime_onehot_or_circular_encoder_1(self.data.index,
                                                                           country=self.country)
            datetime_onehot_or_circular_encoder_2 = DatetimeOnehotORCircularEncoder(to_encoding_args=('month',
                                                                                                      'weekday'),
                                                                                    mode='circular')
            time_var_transformed_2 = datetime_onehot_or_circular_encoder_2(index_source,
                                                                           country=self.country)
            # 其它
            transform_args = self._get_transform_args(transform_args_file_path)
            other_var_transformed = pd.DataFrame(columns=transform_args.columns)
            for this_col in transform_args.columns:
                other_var_transformed.loc[:, this_col] = (self.data.loc[:, this_col] -
                                                          transform_args.loc['minimum', this_col]) / (
                                                                 transform_args.loc['maximum', this_col] -
                                                                 transform_args.loc['minimum', this_col])
            transformed_x_y = pd.concat(
                (time_var_transformed_1, time_var_transformed_2,
                 other_var_transformed.reset_index().drop('TIMESTAMP', axis=1)), axis=1)
            # TODO 泛化输入的维度，sa又要对比考虑/不考虑HEP in the input...好无聊啊
            data_x = torch.tensor(transformed_x_y.iloc[:, :].values,
                                  device='cuda:0',
                                  dtype=torch.float)
            if mode == 'NILM':
                # TODO
                data_y = torch.tensor(transformed_x_y.iloc[:, [-1]].values,
                                      device='cuda:0',
                                      dtype=torch.float)
            elif mode == 'forecast':
                data_y = torch.tensor(transformed_x_y.iloc[:, -2:].values,
                                      device='cuda:0',
                                      dtype=torch.float)
            else:
                raise Exception

            transformed_x_y.index = other_var_transformed.index
            return data_x, data_y, transformed_x_y

        def _get_transform_args(self, file_path: Path) -> pd.DataFrame:
            """
            根据self.weather_var, self.main_var, self.appliance_var，载入 OR 计算并保存min-max变换需要的参数
            """
            if not file_path:
                raise Exception('Should specify file_path')

            @load_exist_pkl_file_otherwise_run_and_save(file_path)
            def func() -> pd.DataFrame:
                transform_args = pd.DataFrame(index=('minimum', 'maximum'),
                                              columns=self.data.columns)
                for i in transform_args.columns:
                    transform_args[i] = (np.nanmin(self.data[i].values), np.nanmax(self.data[i].values))
                return transform_args

            return func()


    class NILMTorchDatasetForecast(NILMTorchDataset):
        def __init__(self, data: pd.DataFrame,
                     *, country,
                     sequence_length: int,
                     over_lapping: Union[bool, int] = False,
                     transform_args_file_path: Path,
                     number_of_previous_days_used_in_input: int = 21):
            super(NILMTorchDatasetForecast, self).__init__(
                data,
                country=country,
                sequence_length=sequence_length,
                over_lapping=over_lapping,
                transform_args_file_path=transform_args_file_path
            )
            self.number_of_previous_days_used_in_input = number_of_previous_days_used_in_input

        def __len__(self):
            return int(self.data.__len__() / self.sequence_length) - self.number_of_previous_days_used_in_input - 1

        def __getitem__(self, index: int):
            # TODO time_variable_dim_number和weather_variable_dim_number泛化
            time_variable_dim_number = 6
            weather_variable_dim_number = 3
            # 决定索引的位置
            # index_slice_x = slice(index * self.sequence_length, (index + 7) * self.sequence_length)
            # index_slice_y = slice((index + 7) * self.sequence_length, (index + 8) * self.sequence_length)
            # TODO 泛化输入的维度，sa又要对比考虑/不考虑HEP in the input...好无聊啊
            data_x = torch.zeros(
                (self.sequence_length,
                 self.transformed_data[0][index].shape[
                     0] + self.number_of_previous_days_used_in_input - 1 + self.number_of_previous_days_used_in_input),
                device='cuda:0'
            )
            # Use d=d0's time information
            index_tuple = (index + self.number_of_previous_days_used_in_input,
                           index + self.number_of_previous_days_used_in_input + 1)
            day_slice = slice(
                index_tuple[0] * self.sequence_length, index_tuple[1] * self.sequence_length
            )
            dim_slice = slice(
                0, time_variable_dim_number
            )
            data_x[:, :time_variable_dim_number] = self.transformed_data[0][day_slice, dim_slice]
            # Use d=(d0-1)'s weather information
            index_tuple = (index + self.number_of_previous_days_used_in_input - 1,
                           index + self.number_of_previous_days_used_in_input)
            day_slice = slice(
                index_tuple[0] * self.sequence_length, index_tuple[1] * self.sequence_length
            )
            dim_slice = slice(
                time_variable_dim_number, time_variable_dim_number + weather_variable_dim_number
            )
            data_x[:, time_variable_dim_number:time_variable_dim_number + weather_variable_dim_number] = \
                self.transformed_data[0][day_slice, dim_slice]
            # Use d=(d0-1)'s, (d0-2)'s, ..., (d0-7)'s total load information
            for i in range(self.number_of_previous_days_used_in_input):
                index_tuple = (index + self.number_of_previous_days_used_in_input - 1 - i,
                               index + self.number_of_previous_days_used_in_input - i)
                day_slice = slice(
                    index_tuple[0] * self.sequence_length, index_tuple[1] * self.sequence_length
                )
                data_x[:, time_variable_dim_number + weather_variable_dim_number + i] = \
                    self.transformed_data[0][day_slice, -2]
                # TODO 泛化输入的维度，sa又要对比考虑/不考虑HEP in the input...好无聊啊
                data_x[:, time_variable_dim_number + weather_variable_dim_number + i + 21] = \
                    self.transformed_data[0][day_slice, -1]

            index_slice_y = slice((index + self.number_of_previous_days_used_in_input) * self.sequence_length,
                                  (index + self.number_of_previous_days_used_in_input + 1) * self.sequence_length)

            data_y = self.transformed_data[1][index_slice_y]  # type: torch.tensor
            return data_x, data_y

        def get_raw_data(self, index):
            index_slice_y = slice((index + self.number_of_previous_days_used_in_input) * self.sequence_length,
                                  (index + self.number_of_previous_days_used_in_input + 1) * self.sequence_length)
            return self.transformed_data[2][slice(
                index * self.sequence_length,
                (index + self.number_of_previous_days_used_in_input) * self.sequence_length
            )], self.transformed_data[2][index_slice_y]

        def _transform(self, transform_args_file_path, mode='forecast') \
                -> Tuple[torch.tensor, torch.tensor, pd.DataFrame]:
            return super(NILMTorchDatasetForecast, self)._transform(transform_args_file_path, mode='forecast')


except ModuleNotFoundError:
    pass

DATASET_ROOT_DIRECTORY = r'E:\OneDrive_Extra\Database\Load_Disaggregation'
LOW_CARBON_LONDON_ROOT = fr'C:\Users\{getpass.getuser()}\OneDrive\PhD\01-PhDProject\Database\UK Power Networks'
TURKEY_ROOT_DIRECTORY = project_path_ / r'Data\Raw\Sasa_Turkey'


def load_low_carbon_london_heat(this_no, df_name: str):
    """
    Depreciated codes. Only for very specific task. Do not be fooled.
    """
    # available_customer_no = (1, 2, 3, 4, 5, 7, 9)
    heat_data_root = Path(
        LOW_CARBON_LONDON_ROOT) / r'Low Carbon London Heat Pump Load Profiles\Data supply 1\Data supply 1\Heat Profiles'
    pq_data_root = Path(
        LOW_CARBON_LONDON_ROOT) / r'Low Carbon London Heat Pump Load Profiles\Data supply 1\Data supply 1\Power Quality'
    figure_buffer_list = []
    names_list = []

    # for this_no in available_customer_no:
    this_figure_buffer_list_elec = []
    this_figure_buffer_list_heat = []
    this_figure_buffer_list_elec_and_heat = []

    # read heat pump
    this_heat_df = pd.read_csv(heat_data_root / f'S1_Customer_L_{this_no}.csv',
                               sep=',', index_col='timestamp')
    this_heat_df = TimeSeries(
        this_heat_df[['external_temperature', 'zone_1_temperature', 'heat_pump_energy_consumption']],
        index=pd.DatetimeIndex(this_heat_df.index)
    )
    this_heat_df['heat_pump_energy_consumption'].iloc[1:] = np.diff(
        this_heat_df['heat_pump_energy_consumption'].values
    )
    this_heat_df.rename({'heat_pump_energy_consumption': 'heat_pump_energy_consumption diff'}, axis=1, inplace=True)
    this_heat_df = this_heat_df.iloc[1:]
    # read electric
    this_pq_df_index = pd.read_csv(pq_data_root / f'S1_Customer_PQ_{this_no}.csv', nrows=2, header=None)
    this_pq_df_index = pd.date_range(start=pd.to_datetime(this_pq_df_index.iloc[0, 1], format="%d/%m/%Y %H:%M"),
                                     end=pd.to_datetime(this_pq_df_index.iloc[1, 1], format="%d/%m/%Y %H:%M"),
                                     freq='T')
    this_pq_df = pd.read_csv(pq_data_root / f'S1_Customer_PQ_{this_no}.csv',
                             sep=',', skiprows=5,
                             engine='python', usecols=['kW of Vln * Il - Mean [kW]'])
    # Apply Sasa's method. He said to take the absolute values for negative power, despite they may take 100 %
    this_pq_df[this_pq_df < 0] = np.abs(this_pq_df[this_pq_df < 0])
    try:
        this_pq_df.index = this_pq_df_index
        this_pq_df_aggregate = this_pq_df.resample('30T').mean()
    except ValueError:
        names_list.append(' ')
        figure_buffer_list.append([])
        return

    # Try align heating and elec
    this_pq_df_and_heat_df = merge_two_time_series_df(this_pq_df,
                                                      this_heat_df,
                                                      interpolate_method='nearest')
    this_pq_df_and_heat_df = this_pq_df_and_heat_df.resample('30T').mean()

    # %% Plot electric all
    def plot_for_original_elec(x, y):
        return time_series(x=x, y=y, color='b',
                           x_label='Date time (original, resolution: 1 minute)', y_label='Active power [kW]',
                           save_to_buffer=True)

    def plot_for_aggregate_elec(x, y):
        return time_series(x=x, y=y, color='b',
                           x_label='Date time (aggregated, resolution: 30 minute)', y_label='Active power [kW]',
                           save_to_buffer=True)

    def plot_for_heat(x, y):
        return time_series(x=x, y=y, color='r', linestyle='--',
                           x_label='Date time (original, resolution: 30 minute)',
                           y_label='heat_pump_energy_consumption\n'
                                   'difference [unknown unit]',
                           save_to_buffer=True)

    def plot_for_aggregate_elec_and_heat(x_elec, y_elec, x_heat, y_heat):
        buffer_1 = plot_for_aggregate_elec(x_elec, y_elec)
        buffer_2 = plot_for_heat(x_heat, y_heat)
        return buffer_1, buffer_2

    this_figure_buffer_list_elec.append(plot_for_original_elec(this_pq_df.index, this_pq_df.iloc[:, 0].values))
    this_figure_buffer_list_elec.append(
        plot_for_aggregate_elec(this_pq_df_aggregate.index, this_pq_df_aggregate.iloc[:, 0].values)
    )

    this_figure_buffer_list_heat.append(plot_for_heat(this_heat_df.index, this_heat_df.iloc[:, 0].values))

    this_figure_buffer_list_elec_and_heat.extend(plot_for_aggregate_elec_and_heat(
        x_elec=this_pq_df_and_heat_df.index,
        y_elec=this_pq_df_and_heat_df.iloc[:, 0],
        x_heat=this_pq_df_and_heat_df.index,
        y_heat=this_pq_df_and_heat_df.iloc[:, -1]
    ))
    iter_dict = {'electric': this_pq_df,
                 'heat': this_heat_df,
                 'electric and heat': this_pq_df_and_heat_df}
    considered_df = iter_dict[df_name]
    figure_buffer_list_specific = copy.deepcopy(figure_buffer_list)
    this_figure_buffer_list_elec_specific = copy.deepcopy(this_figure_buffer_list_elec)
    date_range = pd.date_range(considered_df.first_valid_index().date(),
                               considered_df.last_valid_index().date(),
                               freq='D')
    for i in range(date_range.__len__()):
        if i == date_range.__len__() - 1:
            if considered_df.loc[date_range[i]:].size < 40:
                break
            this_df_i = considered_df.loc[date_range[i]:]
            if df_name == 'electric':
                this_pq_df_aggregate_i = this_pq_df_aggregate.loc[date_range[i]:]

        else:
            this_df_i = considered_df.loc[date_range[i]:date_range[i + 1]].iloc[:-1]
            if df_name == 'electric':
                this_pq_df_aggregate_i = this_pq_df_aggregate.loc[date_range[i]:date_range[i + 1]].iloc[:-1]
        if df_name == 'electric':
            this_figure_buffer_list_elec_specific.append(
                plot_for_original_elec(this_df_i.index, this_df_i.iloc[:, 0].values))
            this_figure_buffer_list_elec_specific.append(
                plot_for_aggregate_elec(this_pq_df_aggregate_i.index, this_pq_df_aggregate_i.iloc[:, 0].values)
            )
            figure_buffer_list_specific.append(this_figure_buffer_list_elec_specific)

        elif df_name == 'heat':
            this_figure_buffer_list_heat.append(
                plot_for_heat(this_df_i.index, this_df_i.iloc[:, -1].values))
            figure_buffer_list_specific.append(this_figure_buffer_list_heat)

        else:
            this_figure_buffer_list_elec_and_heat.extend(
                plot_for_aggregate_elec_and_heat(this_df_i.index, this_df_i.iloc[:, 0].values,
                                                 this_df_i.index, this_df_i.iloc[:, -1].values)
            )
            figure_buffer_list_specific.append(this_figure_buffer_list_elec_and_heat)

    names_list.append(f'S1_Customer_L_{this_no}.csv')
    # Write
    document = docx_document_template_to_collect_figures()
    for i in range(names_list.__len__()):
        document.add_heading(names_list[i], level=1)
        p = document.add_paragraph()
        p.add_run().add_break()
        for j in range(figure_buffer_list_specific[i].__len__()):
            p = document.add_paragraph()
            p.add_run().add_picture(figure_buffer_list_specific[i][j], width=Cm(14))
            # if (j % 2) == 1:
            #     p.add_run().add_break()
        document.add_page_break()
    document.save(f'.\\London {df_name} data household_{this_no}.docx')


def load_ampds2_weather():
    """
    载入ampds2的天气数据
    :return:
    """
    _path = os.path.join(DATASET_ROOT_DIRECTORY,
                         r'AMPds2/dataverse_files/Climate_HourlyWeather.csv')
    reading = pd.read_csv(_path,
                          sep=',')
    read_results = reading[['Temp (C)', 'Rel Hum (%)', 'Stn Press (kPa)']]
    read_results.index = pd.DatetimeIndex(pd.to_datetime(reading.iloc[:, 0],
                                                         format='%Y-%m-%d %H:%M'))
    # _path = os.path.join(DATASET_ROOT_DIRECTORY,
    #                      r'AMPds2/MERRA2/')
    # # America / Vancouver
    # if not os.path.exists(_path + "weather.pkl"):
    #     read_results = pd.DataFrame()
    #     for file in ('2012.csv', '2013.csv', '2014.csv'):
    #         reading = pd.read_csv(_path + file,
    #                               sep=',',
    #                               skiprows=3)
    #         read_results = pd.concat((read_results,
    #                                   pd.DataFrame(index=pd.DatetimeIndex(reading['local_time']),
    #                                                data={'temperature': reading['temperature'].values,
    #                                                      'solar irradiation': reading['radiation_surface'].values,
    #                                                      'precipitation': reading['precipitation'].values,
    #                                                      'air density': reading['air_density'].values})
    #                                   ))
    #     read_results = read_results.loc[~read_results.index.duplicated()]
    #     read_results.to_pickle(_path + "weather.pkl")
    # else:
    #     read_results = pd.read_pickle(_path + "weather.pkl")  # type: pd.DataFrame
    #
    return read_results


def convert_refit_to_h5():
    """
    将refit的csv全部转成一个h5
    :return:
    """
    h5_file_ = os.path.join(DATASET_ROOT_DIRECTORY,
                            r'REFIT\Cleaned\CLEAN_REFIT_081116\Refit.h5')
    if try_to_find_file(h5_file_):
        return
    convert_refit(input_path=os.path.join(DATASET_ROOT_DIRECTORY,
                                          r'REFIT\Cleaned\CLEAN_REFIT_081116'),
                  output_filename=h5_file_,
                  format='HDF')


def load_datasets():
    """
    载入全部三个数据集，h5格式
    :return: 三个数据集的tuple
    """
    _ampds2_dataset = DataSet(os.path.join(DATASET_ROOT_DIRECTORY,
                                           r'AMPds2\dataverse_files\AMPds2.h5'))
    _refit_dataset = DataSet(os.path.join(DATASET_ROOT_DIRECTORY,
                                          r'REFIT\Cleaned\CLEAN_REFIT_081116\Refit.h5'))

    convert_refit_to_h5()
    _uk_dale_dataset = DataSet(os.path.join(DATASET_ROOT_DIRECTORY,
                                            r'UK-DALE\ukdale.h5\ukdale.h5'))
    return _ampds2_dataset, _refit_dataset, _uk_dale_dataset


def get_training_set_and_test_set_for_ampds2_dataset():
    """
    从ampds2_dataset中分离出training set和test set，
    只考虑building 1
    2012-4-1 0:00到2013-4-1 0:00 是training set
    2013-4-1 0:00到2014-4-1 0:00 是test set
    :return:
    """
    training_set, _, _ = load_datasets()
    test_set, _, _ = load_datasets()
    whole_set, _, _ = load_datasets()

    training_set.set_window(end='2013-4-1')
    training_set = training_set.buildings[1].elec

    test_set.set_window(start='2013-4-1')
    test_set = test_set.buildings[1].elec

    whole_set = whole_set.buildings[1].elec
    return training_set, test_set, whole_set


def ampds2_dataset_full_df(resolution: int) -> pd.DataFrame:
    """
    ampds2_dataset的heat，main，和对应的气象，和对应的时间
    """

    @load_exist_pkl_file_otherwise_run_and_save(project_path_ / r'Data\Raw\for_Energies_Research_paper_2020\Ampds2.pkl')
    def func():
        _, _, ampds2 = get_training_set_and_test_set_for_ampds2_dataset()
        heating_df = next(ampds2.select_using_appliances(
            original_name='HPE').meters[0].load(ac_type='active', sample_period=resolution))
        heating_df = heating_df.droplevel('physical_quantity', axis=1)  # type: DataFrame
        heating_df.rename(columns={'active': 'HPE'}, inplace=True)

        mains_df = next(ampds2.mains().load(
            ac_type='active', sample_period=resolution)).droplevel('physical_quantity', axis=1)  # type: DataFrame
        mains_df.rename(columns={mains_df.columns[0]: 'Mains'}, inplace=True)

        ampds2_weather_df = load_ampds2_weather()
        mains_weather_df_merged = merge_two_time_series_df(mains_df, ampds2_weather_df)
        mains_weather_df_merged = mains_weather_df_merged.reindex(columns=mains_weather_df_merged.columns[1:].append(
            mains_weather_df_merged.columns[slice(1)]))

        full_data_df = pd.concat((mains_weather_df_merged, heating_df), axis=1)

        # WARNING: full_data_df.index is not used! because stupid Pandas force to have DST transformation, causing two
        # 1.00 am problem, despite that the original datetime has already been taken care of!
        # full_data_df['index_index'] = range(full_data_df.shape[0]) # This line is for debug
        # index_source = full_data_df.index # This line is to use full_data_df.index
        index_source = pd.DatetimeIndex(pd.date_range(
            start=datetime.datetime(year=full_data_df.index[0].year,
                                    month=full_data_df.index[0].month,
                                    day=full_data_df.index[0].day,
                                    hour=full_data_df.index[0].hour,
                                    minute=full_data_df.index[0].minute),
            end=datetime.datetime(year=full_data_df.index[-1].year,
                                  month=full_data_df.index[-1].month,
                                  day=full_data_df.index[-1].day,
                                  hour=full_data_df.index[-1].hour,
                                  minute=full_data_df.index[-1].minute),
            freq=full_data_df.index.freq)
        )
        full_data_df['year'] = index_source.year
        full_data_df['month'] = index_source.month
        full_data_df['day'] = index_source.day
        full_data_df['dayofweek'] = index_source.dayofweek + 1
        full_data_df['hour'] = index_source.hour
        full_data_df['minute'] = index_source.minute
        full_data_df['moon_phase'] = moon_phase(full_data_df.index.to_numpy())  # Moon phase must use full_data_df.index
        date_time_one_hot_encoder = DatetimeOnehotORCircularEncoder(to_encoding_args=('holiday', 'summer_time'))
        time_var_transformed = date_time_one_hot_encoder(
            # 'holiday' and 'summer_time' must use full_data_df.index
            full_data_df.index,
            country=Canada()
        )
        full_data_df['holiday'] = np.array(time_var_transformed.iloc[:, 0] == 1, dtype=int)
        full_data_df['summer_time'] = np.array(time_var_transformed.iloc[:, 1] == 1, dtype=int)
        return full_data_df

    return func()


class ScotlandDataset(metaclass=ABCMeta):
    data_path_root = project_path_ / r'Data\Raw\Scotland selected'
    __slots__ = ('name', 'matlab_mat_file_folder', 'dataset')

    def __init__(self, name: str):
        """
        name可选: ['Drum', 'John', 'MAYB', 'STHA']
        """
        if name not in ['Drum', 'John', 'MAYB', 'STHA']:
            raise Exception('Unknown bus')
        self.name = name
        self.matlab_mat_file_folder = self.data_path_root / self.name
        self.dataset = self.load_raw_data()  # type: pd.DataFrame

    def __str__(self):
        return f'Scotland dataset: {self.name}, from {self.dataset.index[0]} to {self.dataset.index[-1]}'

    def load_raw_data(self) -> pd.DataFrame:
        """
        重复利用以前的结果。载入matlab那些原始数据
        """
        time_index = self._get_time_index()
        holiday_ndarray = self._get_holiday_ndarray()
        bst_ndarray = self._get_british_summer_time()
        raw_data = pd.DataFrame(data={'holiday': holiday_ndarray,
                                      'BST': bst_ndarray,
                                      'active power': self.load_active_power_mat(),
                                      'temperature': self.get_temperature()},
                                index=time_index)
        return raw_data

    def set_weekends_and_holiday_to_zeros(self, inplace=False) -> Tuple[pd.DataFrame, ndarray]:
        """
        把节假日或者周末的数据置为0
        :return 置0后的pd.DataFrame和对应的mask
        """
        mask = np.any((self.dataset.index.weekday == 5,  # 周六，因为Monday=0, Sunday=6.
                       self.dataset.index.weekday == 6,  # 周天，因为Monday=0, Sunday=6.
                       self.dataset['holiday'] == 1), axis=0)
        if inplace:
            self.dataset.loc[mask, 'active power'] = 0
            return self.dataset, mask
        else:
            dataset_copy = copy.deepcopy(self.dataset)
            dataset_copy.loc[mask, 'active power'] = 0
            return dataset_copy, mask

    def _get_time_index(self) -> pd.DatetimeIndex:
        """
        基于matlab那些原始数据得到python的pd.DatetimeIndex
        """
        ts_matrix = self.load_ts_mat()
        datetime_tuple = [datetime.datetime(year=int(x[0]),
                                            month=int(x[1]),
                                            day=int(x[2]),
                                            hour=int(x[3]),
                                            minute=int(60 * (x[3] - int(x[3])))) for x in ts_matrix]
        time_index = pd.DatetimeIndex(datetime_tuple)
        return time_index

    def _get_holiday_ndarray(self) -> ndarray:
        """
        基于matlab那些原始数据，提取节假日标志量：1代表是holiday，0代表不是
        """
        ts_matrix = self.load_ts_mat()
        return ts_matrix[:, -1].astype(int)

    def _get_british_summer_time(self) -> ndarray:
        ts_matrix = self.load_ts_mat()
        return ts_matrix[:, -2].astype(int)

    @abstractmethod
    def load_active_power_mat(self) -> ndarray:
        """
        载入Data_P_modified.mat或者Data_P.mat
        """
        pass

    @abstractmethod
    def load_ts_mat(self) -> ndarray:
        """
        载入Data_ts_modified.mat或者Data_ts.mat
        """
        pass

    @abstractmethod
    def get_temperature(self) -> ndarray:
        """
        载入Data_temperature_modified.mat或者Data_temperature.mat
        """
        pass


class ScotlandLongerDataset(ScotlandDataset):
    """
    指那些有四五年记录的bus，比如'Drum'和'John'
    """

    def __init__(self, name: str):
        """
        需要注意的是闰年2月29号没有记录，active power全部用0去填充的
        """
        if name not in ('Drum', 'John'):
            raise Exception('Wrong name')
        super().__init__(name)

    def load_active_power_mat(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_P_modified.mat')['P'].flatten()

    def load_ts_mat(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_ts_modified.mat')['ts']

    def get_temperature(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_temperature_modified.mat')['temperature'].flatten()


class ScotlandShorterDataset(ScotlandDataset):
    """
    指那些只有一年记录的bus，比如'MAYB'和'STHA'
    """

    def __init__(self, name: str):
        if name not in ('MAYB', 'STHA'):
            raise Exception('Wrong name')
        super().__init__(name)

    def load_active_power_mat(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_P.mat')['P'].flatten()

    def load_ts_mat(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_ts.mat')['ts']

    def get_temperature(self) -> ndarray:
        return loadmat(self.matlab_mat_file_folder / 'Data_temperature.mat')['temperature'].flatten()


def load_turkey_dataset():
    @load_exist_pkl_file_otherwise_run_and_save(project_path_ / r'Data\Raw\for_Energies_Research_paper_2020\Turkey.pkl')
    def func():
        turkey_data = {}
        for file_prefix in ('Apartment', 'Detached House'):
            # %% read mains and lighting
            one_excel_reading = pd.read_excel(TURKEY_ROOT_DIRECTORY / f'{file_prefix} Meter Data.xlsx',
                                              sheet_name=None)  # type: dict

            reading_df = pd.DataFrame()
            for i, (key, val) in enumerate(one_excel_reading.items()):
                if ('Main' not in key) and ('Lighting' not in key):
                    continue
                datetime_index = pd.DatetimeIndex(pd.to_datetime(val.values[:, 0]))
                this_sheet_df = pd.DataFrame(data=val['Active'].values,
                                             index=datetime_index.round('H'),
                                             columns=['main' if 'Main' in key else 'lighting']).sort_index()
                this_sheet_df = this_sheet_df[~this_sheet_df.index.duplicated(keep='first')]
                if i == 0:
                    reading_df = this_sheet_df
                else:
                    reading_df = pd.merge(reading_df, this_sheet_df, left_index=True, right_index=True, how='outer')
            reading_df = reading_df.diff()

            # %% Weather
            weather_df = pd.DataFrame()
            for i, year in enumerate((2017, 2018)):
                this_weather_df = pd.read_csv(TURKEY_ROOT_DIRECTORY / f'{file_prefix}_MERRA_II_{year}.csv',
                                              skiprows=3)
                this_weather_df.index = pd.DatetimeIndex(pd.to_datetime(this_weather_df['local_time'].values))
                this_weather_df = this_weather_df.iloc[:, 2:]
                if i == 0:
                    weather_df = this_weather_df
                else:
                    weather_df = pd.concat((weather_df, this_weather_df))  # type: pd.DataFrame

            # %% Merge
            merged_reading = pd.merge(reading_df, weather_df, left_index=True, right_index=True, how='left')
            #
            turkey_data[file_prefix] = merged_reading
        return turkey_data

    return func()


def pre_call_():
    """
    This is to save the followings to pd.DataFrame format:
    - Ampds 2, main, heat, weather
    - UKDale
    Please call this function in NILM_Project conda env, as Python_Project conda env has very latest packages that
    nilmtk does not support!
    In the actual analysis, Python_Project should be used.
    The only reason NILM_Project exists is to transform the required data (especially for Energies_Research_paper_2020)
    to pd.DataFrame.
    """
    sample_period = 3600
    # %% Ampds2
    ampds2_dataset_full_df(sample_period)

    # %% UK Dale
    @load_exist_pkl_file_otherwise_run_and_save(project_path_ / r'Data\Raw\for_Energies_Research_paper_2020\UKDALE.pkl')
    def func():
        _, _, uk_dale = load_datasets()

        uk_dale_lighting = next(uk_dale.buildings[1].elec.select_using_appliances(category='lighting').load(
            sample_period=sample_period))[('power', 'active')]
        uk_dale_heating = next(uk_dale.buildings[1].elec.select_using_appliances(category='heating').load(
            sample_period=sample_period))[('power', 'active')]
        uk_dale_mains = uk_dale.buildings[1].elec.mains().power_series_all_data(sample_period=sample_period)

        uk_dale_df = pd.DataFrame()
        names = ('lighting', 'heating', 'mains')
        for i, this_df in enumerate((uk_dale_lighting, uk_dale_heating, uk_dale_mains)):
            this_df = this_df[~this_df.index.duplicated(keep='first')]
            this_df = pd.DataFrame(data=this_df.values,
                                   index=this_df.index,
                                   columns=[names[i]])
            if i == 0:
                uk_dale_df = this_df
            else:
                uk_dale_df = pd.merge(uk_dale_df,
                                      this_df, left_index=True, right_index=True, how='outer')
        return uk_dale_df

    func()


if __name__ == '__main__':
    # ampds2_dataset, refit_dataset, uk_dale_dataset = load_datasets()
    # get_training_set_and_test_set_for_ampds2_dataset()
    # John_data = ScotlandLongerDataset('John')
    # John_data.set_weekends_and_holiday_to_zeros()
    # for _this_no in (1, 2, 3, 4, 5, 7):
    #     for this_type in ('heat', 'electric', 'electric and heat'):
    #         load_low_carbon_london_heat(_this_no, this_type)
    # load_ampds2_weather()
    # ampds2_dataset_full_df(30 * 60)

    # %% Execute in NILM_Project, preparing
    pre_call_()

    # %% Test codes
